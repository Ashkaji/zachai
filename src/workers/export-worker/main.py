import base64
import hashlib
import io
import json
import logging
import os
import uuid
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from fastapi import FastAPI, Header, HTTPException
from minio import Minio
from minio.error import S3Error
from pydantic import BaseModel, Field, field_validator

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("export-worker")

app = FastAPI(title="export-worker")

required_env = ["MINIO_ENDPOINT", "MINIO_ACCESS_KEY", "MINIO_SECRET_KEY"]
missing = [env for env in required_env if env not in os.environ]
if missing:
    raise RuntimeError(f"Missing required environment variables: {', '.join(missing)}")

SNAPSHOT_BUCKET = (os.environ.get("SNAPSHOT_BUCKET") or "snapshots").strip() or "snapshots"
SNAPSHOT_CALLBACK_SECRET = (os.environ.get("SNAPSHOT_CALLBACK_SECRET") or "").strip()
DLQ_DIR = Path(os.environ.get("EXPORT_WORKER_DLQ_DIR") or "/tmp/export-worker-dlq")
MAX_RETRIES = int(os.environ.get("EXPORT_WORKER_MAX_RETRIES") or "3")

minio_client = Minio(
    endpoint=os.environ["MINIO_ENDPOINT"],
    access_key=os.environ["MINIO_ACCESS_KEY"],
    secret_key=os.environ["MINIO_SECRET_KEY"],
    secure=os.environ.get("MINIO_SECURE", "false").lower() == "true",
)


class SnapshotExportRequest(BaseModel):
    document_id: int = Field(..., gt=0)
    yjs_state_binary: str = Field(..., min_length=1, max_length=8_000_000)

    @field_validator("yjs_state_binary")
    @classmethod
    def _validate_b64(cls, v: str) -> str:
        try:
            base64.b64decode(v.encode("ascii"), validate=True)
        except Exception as exc:
            raise ValueError("yjs_state_binary must be valid base64") from exc
        return v


def _require_secret(x_secret: str | None) -> None:
    if not SNAPSHOT_CALLBACK_SECRET:
        raise HTTPException(status_code=503, detail={"error": "Snapshot secret not configured"})
    if not x_secret:
        raise HTTPException(status_code=401, detail={"error": "Unauthorized"})
    if x_secret != SNAPSHOT_CALLBACK_SECRET:
        raise HTTPException(status_code=403, detail={"error": "Forbidden"})


def _sha256_hex(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _upload_and_verify_object(object_name: str, data: bytes, content_type: str) -> str:
    checksum = _sha256_hex(data)
    minio_client.put_object(
        bucket_name=SNAPSHOT_BUCKET,
        object_name=object_name,
        data=io.BytesIO(data),
        length=len(data),
        content_type=content_type,
        metadata={"sha256": checksum},
    )
    stat = minio_client.stat_object(SNAPSHOT_BUCKET, object_name)
    md = {str(k).lower(): str(v) for k, v in (stat.metadata or {}).items()}
    persisted = md.get("x-amz-meta-sha256") or md.get("sha256")
    if persisted != checksum:
        raise RuntimeError(f"checksum mismatch for {object_name}")
    return checksum


def _build_docx_bytes(document_id: int, snapshot_id: str, yjs_raw: bytes) -> bytes:
    doc = Document()
    doc.add_heading(f"ZachAI Snapshot {snapshot_id}", 1)
    doc.add_paragraph(f"Document ID: {document_id}")
    doc.add_paragraph(f"Created at (UTC): {datetime.now(timezone.utc).isoformat()}")
    doc.add_paragraph("Note: This DOCX snapshot stores metadata and Yjs state checksum for restore workflows.")
    doc.add_paragraph(f"Yjs bytes length: {len(yjs_raw)}")
    doc.add_paragraph(f"Yjs SHA-256: {_sha256_hex(yjs_raw)}")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _write_dlq(payload: dict, reason: str) -> Path:
    DLQ_DIR.mkdir(parents=True, exist_ok=True)
    name = f"{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}-{uuid.uuid4().hex[:8]}.json"
    path = DLQ_DIR / name
    path.write_text(json.dumps({"reason": reason, "payload": payload}, ensure_ascii=True), encoding="utf-8")
    return path


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/snapshot-export")
def snapshot_export(
    body: SnapshotExportRequest,
    x_snapshot_secret: str | None = Header(default=None, alias="X-ZachAI-Snapshot-Secret"),
) -> dict:
    _require_secret(x_snapshot_secret)

    yjs_raw = base64.b64decode(body.yjs_state_binary.encode("ascii"), validate=True)
    yjs_sha = _sha256_hex(yjs_raw)
    ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    snapshot_id = f"{ts}-{uuid.uuid4().hex[:10]}"
    prefix = f"{body.document_id}/{snapshot_id}"
    json_key = f"{prefix}.json"
    docx_key = f"{prefix}.docx"

    json_payload = {
        "snapshot_id": snapshot_id,
        "document_id": body.document_id,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "yjs_state_binary": body.yjs_state_binary,
        "yjs_sha256": yjs_sha,
        "source": "hocuspocus-idle-callback",
    }
    json_bytes = json.dumps(json_payload, ensure_ascii=True, separators=(",", ":")).encode("utf-8")
    docx_bytes = _build_docx_bytes(body.document_id, snapshot_id, yjs_raw)

    last_error: Exception | None = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            json_sha = _upload_and_verify_object(json_key, json_bytes, "application/json")
            docx_sha = _upload_and_verify_object(
                docx_key,
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            logger.info(
                "snapshot_export_ok document_id=%s snapshot_id=%s attempt=%s json_key=%s docx_key=%s",
                body.document_id,
                snapshot_id,
                attempt,
                json_key,
                docx_key,
            )
            return {
                "status": "ok",
                "snapshot_id": snapshot_id,
                "bucket": SNAPSHOT_BUCKET,
                "json_object_key": f"{SNAPSHOT_BUCKET}/{json_key}",
                "docx_object_key": f"{SNAPSHOT_BUCKET}/{docx_key}",
                "yjs_sha256": yjs_sha,
                "json_sha256": json_sha,
                "docx_sha256": docx_sha,
                "attempts": attempt,
            }
        except (S3Error, RuntimeError, OSError) as exc:
            last_error = exc
            logger.warning(
                "snapshot_export_retry document_id=%s snapshot_id=%s attempt=%s/%s error=%s",
                body.document_id,
                snapshot_id,
                attempt,
                MAX_RETRIES,
                exc,
            )

    dlq_path = _write_dlq(json_payload, str(last_error) if last_error else "unknown error")
    logger.error(
        "snapshot_export_dlq document_id=%s snapshot_id=%s dlq=%s error=%s",
        body.document_id,
        snapshot_id,
        dlq_path,
        last_error,
    )
    raise HTTPException(
        status_code=500,
        detail={
            "error": "Snapshot export failed after retries",
            "snapshot_id": snapshot_id,
            "dlq_path": str(dlq_path),
        },
    )
